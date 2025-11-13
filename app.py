import os, io, re, tempfile
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from werkzeug.utils import secure_filename
import pytesseract
from PIL import Image
from pdf2image import convert_from_bytes
import mammoth
from docx import Document
from openpyxl import Workbook
from datetime import datetime

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)

ALLOWED_EXT = {'pdf','docx','png','jpg','jpeg'}
UPLOAD_DIR = '/tmp/uyviet_uploads'
os.makedirs(UPLOAD_DIR, exist_ok=True)

def allowed(filename):
    return '.' in filename and filename.rsplit('.',1)[1].lower() in ALLOWED_EXT

def ocr_image_bytes(image_bytes):
    img = Image.open(io.BytesIO(image_bytes)).convert('RGB')
    text = pytesseract.image_to_string(img, lang='vie+eng')
    return text

def extract_docx_text(file_bytes):
    result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
    return result.value

def extract_pdf_text(file_bytes):
    text_total = []
    # convert to images (one image per page)
    images = convert_from_bytes(file_bytes)
    for i, img in enumerate(images, start=1):
        byte_arr = io.BytesIO()
        img.save(byte_arr, format='PNG')
        txt = pytesseract.image_to_string(Image.open(io.BytesIO(byte_arr.getvalue())), lang='vie+eng')
        text_total.append(f"--- PAGE {i} ---\n{txt}")
    return "\n\n".join(text_total)

def parse_project_fields(text):
    out = {}
    lower = text.lower()
    def find_after(keys):
        for k in keys:
            idx = lower.find(k)
            if idx>=0:
                snippet = text[idx: idx+300].splitlines()
                for s in snippet:
                    s = s.strip()
                    if s and len(s) > len(k):
                        # remove key from line
                        s_clean = re.sub(re.escape(k), '', s, flags=re.IGNORECASE).strip(": -–—")
                        s_clean = re.sub(r'\s{2,}', ' ', s_clean).strip()
                        if s_clean:
                            return s_clean
        return ""
    out['project_name'] = find_after(['tên dự án','dự án','project:','project'])
    out['client'] = find_after(['chủ đầu tư','chủ đầu tư:','owner:','client:'])
    out['location'] = find_after(['địa điểm','vị trí','location:'])
    out['scale'] = find_after(['quy mô','diện tích','scale:','area:'])
    out['functions'] = find_after(['công năng','chức năng','functions:'])
    out['special'] = find_after(['pccc','phòng cháy','cách âm','tiết kiệm năng lượng','yêu cầu đặc biệt'])
    return out

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/upload', methods=['POST'])
def api_upload():
    f = request.files.get('file')
    if not f or not allowed(f.filename):
        return jsonify({'error':'No file or file type not allowed'}), 400
    filename = secure_filename(f.filename)
    data = f.read()
    ext = filename.rsplit('.',1)[1].lower()
    extracted = ""
    try:
        if ext == 'docx':
            extracted = extract_docx_text(data)
        elif ext == 'pdf':
            extracted = extract_pdf_text(data)
        else:
            extracted = ocr_image_bytes(data)
    except Exception as e:
        return jsonify({'error':str(e)}), 500
    fields = parse_project_fields(extracted)
    return jsonify({'text': extracted, 'fields': fields})

@app.route('/api/generate', methods=['POST'])
def api_generate():
    payload = request.json or {}
    content = payload.get('content','')
    meta = payload.get('meta',{})
    project_short = meta.get('project_short') or re.sub(r'[^A-Za-z0-9]+','_', meta.get('project_name','Project')).strip('_')[:30]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    base = f"UyViet_{project_short}_GiaiPhap_PromptV4_{timestamp}"
    # create docx
    doc = Document()
    # Title and header
    doc.add_heading("CÔNG TY CỔ PHẦN TƯ VẤN ĐẦU TƯ XÂY DỰNG UY VIỆT", level=1)
    doc.add_heading("HỒ SƠ: GIẢI PHÁP & PHƯƠNG PHÁP LUẬN", level=2)
    doc.add_paragraph(f"Tên dự án: {meta.get('project_name','')}")
    doc.add_paragraph(f"Chủ đầu tư: {meta.get('client','')}")
    doc.add_paragraph(f"Địa điểm: {meta.get('location','')}")
    doc.add_paragraph("")
    # Add content paragraphs
    for line in content.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    docx_name = base + ".docx"
    return send_file(docx_io, as_attachment=True, download_name=docx_name, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/api/download_tools', methods=['GET'])
def api_tools():
    wb = Workbook()
    ws = wb.active
    ws.title = "RACI Matrix"
    ws.append(["Hạng mục","R","A","C","I"])
    ws.append(["Thiết kế kiến trúc","KTS","TGĐ","Điều phối","CĐT"])
    ws.append(["Thiết kế kết cấu","KS KC","TGĐ","Điều phối","CĐT"])
    ws.append(["Thiết kế M&E","KS M&E","TGĐ","Điều phối","CĐT"])
    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name="BoMau_VisualDashboards_UyViet.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT',8080)))
